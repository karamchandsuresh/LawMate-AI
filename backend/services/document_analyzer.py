import os
from io import BytesIO

from docx import Document
from PIL import Image
from pypdf import PdfReader
import pytesseract
pytesseract.pytesseract.tesseract_cmd = (
    r"C:\Program Files\Tesseract-OCR\tesseract.exe"
)

from dotenv import load_dotenv
from google import genai


# ============================================================
# ENVIRONMENT
# ============================================================

load_dotenv()

GEMINI_API_KEY = os.getenv(
    "GEMINI_API_KEY"
)

if not GEMINI_API_KEY:
    raise ValueError(
        "GEMINI_API_KEY not found in backend/.env"
    )


# ============================================================
# GEMINI
# ============================================================

GEMINI_MODEL = "gemini-3.1-flash-lite"

gemini_client = genai.Client(
    api_key=GEMINI_API_KEY
)


# ============================================================
# SUPPORTED FILE TYPES
# ============================================================

SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".jpg",
    ".jpeg",
    ".png",
}


# ============================================================
# TEXT NORMALIZATION
# ============================================================

def normalize_text(text):
    """
    Normalize extracted document text.
    """

    if not text:
        return ""

    lines = []

    for line in text.splitlines():

        line = " ".join(
            line.strip().split()
        )

        if line:
            lines.append(
                line
            )

    return "\n".join(
        lines
    ).strip()


# ============================================================
# PDF EXTRACTION
# ============================================================

def extract_pdf_text(
    file_bytes
):
    """
    Extract text from a text-based PDF.
    """

    reader = PdfReader(
        BytesIO(file_bytes)
    )

    pages = []

    for page_number, page in enumerate(
        reader.pages,
        start=1
    ):

        try:

            page_text = (
                page.extract_text()
            )

        except Exception as error:

            print(
                f"PDF extraction warning "
                f"on page {page_number}: "
                f"{error}"
            )

            continue

        if page_text:

            page_text = normalize_text(
                page_text
            )

            if page_text:

                pages.append(
                    f"--- PAGE {page_number} ---\n"
                    f"{page_text}"
                )

    return "\n\n".join(
        pages
    ).strip()


# ============================================================
# DOCX EXTRACTION
# ============================================================

def extract_docx_text(
    file_bytes
):
    """
    Extract text from Microsoft Word DOCX files.
    """

    document = Document(
        BytesIO(file_bytes)
    )

    parts = []

    # Paragraphs
    for paragraph in document.paragraphs:

        text = normalize_text(
            paragraph.text
        )

        if text:
            parts.append(
                text
            )

    # Tables
    for table in document.tables:

        for row in table.rows:

            row_text = []

            for cell in row.cells:

                cell_text = normalize_text(
                    cell.text
                )

                if cell_text:
                    row_text.append(
                        cell_text
                    )

            if row_text:

                parts.append(
                    " | ".join(
                        row_text
                    )
                )

    return "\n".join(
        parts
    ).strip()


# ============================================================
# IMAGE OCR
# ============================================================

def extract_image_text(
    file_bytes
):
    """
    Extract text from JPG/JPEG/PNG using Tesseract OCR.
    """

    image = Image.open(
        BytesIO(file_bytes)
    )

    image = image.convert(
        "RGB"
    )

    try:

        text = pytesseract.image_to_string(
            image,
            lang="eng"
        )

    except pytesseract.TesseractNotFoundError:

        raise RuntimeError(
            "Tesseract OCR engine is not installed "
            "or is not available in PATH."
        )

    return normalize_text(
        text
    )


# ============================================================
# FILE TEXT EXTRACTION ROUTER
# ============================================================

def extract_text_from_file(
    filename,
    file_bytes
):
    """
    Select extraction method based on file extension.
    """

    extension = os.path.splitext(
        filename.lower()
    )[1]

    if extension not in SUPPORTED_EXTENSIONS:

        raise ValueError(
            "Unsupported file type. "
            "Supported formats: PDF, DOCX, JPG, JPEG, PNG."
        )

    if extension == ".pdf":

        return extract_pdf_text(
            file_bytes
        )

    if extension == ".docx":

        return extract_docx_text(
            file_bytes
        )

    if extension in {
        ".jpg",
        ".jpeg",
        ".png",
    }:

        return extract_image_text(
            file_bytes
        )

    raise ValueError(
        "Unsupported document format."
    )


# ============================================================
# ANALYZE LEGAL DOCUMENT
# ============================================================

def analyze_legal_document(
    filename,
    extracted_text
):
    """
    Analyze extracted legal-document text with Gemini.
    """

    if not extracted_text.strip():

        raise ValueError(
            "No readable text was extracted "
            "from the uploaded document."
        )

    # Avoid sending extremely large files in one request.
    #
    # For Version 1 we analyze the first 30,000 characters.
    # Later we can upgrade this to chunked document analysis.

    analysis_text = (
        extracted_text[:30000]
    )

    prompt = f"""
You are LawMate AI, an Indian legal document analysis assistant.

Analyze the uploaded document carefully.

The document may be:

- a legal notice
- agreement
- contract
- complaint
- court document
- government document
- affidavit
- petition
- employment document
- property document
- consumer document
- or another legal document

Use ONLY the uploaded document text when describing
specific contents of the document.

Do not invent clauses, parties, dates, obligations,
penalties, or legal provisions.

If information is unclear or missing, clearly say so.

============================================================
DOCUMENT
============================================================

Filename:
{filename}

Extracted Text:

{analysis_text}

============================================================
OUTPUT FORMAT
============================================================

📄 Document Overview

Explain what type of document this appears to be and
its main purpose.

📝 Summary

Provide a concise summary of the document.

👥 Parties / Important Entities

Identify people, companies, authorities, courts,
institutions, or other important entities mentioned.

📌 Key Clauses / Legal Points

List the important clauses, obligations, rights,
conditions, deadlines, or legal provisions found
in the document.

⚠️ Potential Issues / Risks

Identify potentially important legal risks,
ambiguous clauses, unusual obligations, missing
information, or matters that deserve closer review.

✅ Suggested Next Steps

Provide practical next steps the user may consider.

If professional legal review may be important,
say so clearly.

⚠️ Disclaimer

State that LawMate AI provides document information
for educational assistance and does not replace
professional legal advice.
"""

    response = (
        gemini_client.models.generate_content(
            model=GEMINI_MODEL,
            contents=prompt
        )
    )

    return response.text


# ============================================================
# COMPLETE DOCUMENT ANALYZER
# ============================================================

def process_document(
    filename,
    file_bytes
):
    """
    Complete pipeline:

    uploaded file
        ↓
    extract text
        ↓
    Gemini analysis
        ↓
    structured result
    """

    extracted_text = (
        extract_text_from_file(
            filename,
            file_bytes
        )
    )

    if len(
        extracted_text.strip()
    ) < 20:

        raise ValueError(
            "The uploaded document contains too little "
            "readable text for analysis."
        )

    analysis = (
        analyze_legal_document(
            filename,
            extracted_text
        )
    )

    return {
        "filename": filename,
        "characters_extracted": len(
            extracted_text
        ),
        "analysis": analysis,
    }


# ============================================================
# LOCAL INFORMATION
# ============================================================

if __name__ == "__main__":

    print(
        "LawMate Document Analyzer service ready."
    )

    print(
        "Supported formats: "
        "PDF, DOCX, JPG, JPEG, PNG"
    )